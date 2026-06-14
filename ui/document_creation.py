import os
import json

from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QPushButton, QHBoxLayout,
    QMessageBox, QLineEdit, QScrollArea, QComboBox, QTextEdit, QApplication, QStyle,
    QTableWidget, QTableWidgetItem, QCheckBox, QFrame, QInputDialog, QFileDialog, QHeaderView,
    QColorDialog, QFontComboBox, QSpinBox
)
from PySide6.QtCore import Qt, QTimer, QEvent
from PySide6.QtGui import QIcon, QTextCursor, QColor, QTextCharFormat
from docx.shared import RGBColor
from docx import Document
from ui.shared_state import set_template
import re


class DocumentCreationScreen(QWidget):
    """Allows the user to type a new document from scratch.

    The screen provides toolbar buttons for adding structured sections,
    bullet lists, or editable tables. Each click inserts a dynamic input field:
    sections consist of a title and paragraph, bullets create numbered entries,
    and tables are editable QTableWidgets with row/column controls. A regular
    free-text editor remains below the dynamic area for miscellaneous content.

    When the user proceeds, all of the dynamic fields (followed by any free
    text) are converted to a python-docx Document object and stored in shared
    state so the rest of the workflow can operate on it.
    """

    @staticmethod
    def get_icon_up():
        return QApplication.style().standardIcon(QStyle.SP_ArrowUp)

    @staticmethod
    def get_icon_down():
        return QApplication.style().standardIcon(QStyle.SP_ArrowDown)

    @staticmethod
    def get_icon_remove():
        return QApplication.style().standardIcon(QStyle.SP_DialogCloseButton)

    # kept for backwards compatibility though no longer inserted by default
    class SectionWidget(QWidget):
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            hl = QHBoxLayout(self)
            self.title = QComboBox()
            self.title.setEditable(True)
            self.title.setInsertPolicy(QComboBox.NoInsert)
            self.title.setCurrentText("")
            self.body = QTextEdit()
            self.body.setPlaceholderText("Section Content")
            self.body.setMaximumHeight(80)
            hl.addWidget(self.title)
            hl.addWidget(self.body)
            def style_btn(btn, text, color="#94a3b8"):
                btn.setText(text)
                btn.setFixedSize(28, 28)
                btn.setStyleSheet(f"""
                    QPushButton {{
                        background-color: #1e293b;
                        color: {color};
                        border: 1px solid #334155;
                        border-radius: 4px;
                        font-weight: bold;
                    }}
                    QPushButton:hover {{
                        background-color: #334155;
                        color: white;
                    }}
                """)

            # Up button
            self.btn_up = QPushButton()
            style_btn(self.btn_up, "🔼")
            hl.addWidget(self.btn_up)
            if move_up_callback:
                self.btn_up.clicked.connect(lambda: move_up_callback(self))
            # Down button
            self.btn_down = QPushButton()
            style_btn(self.btn_down, "🔽")
            hl.addWidget(self.btn_down)
            if move_down_callback:
                self.btn_down.clicked.connect(lambda: move_down_callback(self))
            # Remove button
            self.remove_btn = QPushButton()
            style_btn(self.remove_btn, "✖", "#ef4444")
            hl.addWidget(self.remove_btn)
            if remove_callback:
                self.remove_btn.clicked.connect(lambda: remove_callback(self))

    class HeaderWidget(QWidget):
        """Header entry with H1/H2/H3 level selector."""
        LEVEL_STYLES = {
            1: ("H1", "font-size: 22px; font-weight: bold; color: #f1f5f9;", "Heading 1"),
            2: ("H2", "font-size: 18px; font-weight: bold; color: #cbd5e1;", "Heading 2"),
            3: ("H3", "font-size: 14px; font-weight: bold; color: #94a3b8;", "Heading 3"),
        }

        def __init__(self, level=1, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            self.level = level
            vl = QVBoxLayout(self)
            vl.setContentsMargins(0, 0, 0, 0)
            vl.setSpacing(2)

            # Level selector row
            level_row = QHBoxLayout()
            level_row.setContentsMargins(0, 0, 0, 0)
            level_row.setSpacing(4)
            for lvl in (1, 2, 3):
                btn = QPushButton(f"H{lvl}")
                btn.setFixedSize(32, 22)
                btn.setCheckable(True)
                btn.setChecked(lvl == level)
                btn.setStyleSheet("""
                    QPushButton { background-color: #1e293b; color: #94a3b8; border: 1px solid #334155; border-radius: 3px; font-size: 11px; font-weight: bold; }
                    QPushButton:checked { background-color: #2563eb; color: white; border: 1px solid #3b82f6; }
                    QPushButton:hover { background-color: #334155; }
                """)
                btn.clicked.connect(lambda _, l=lvl, b=btn: self._set_level(l))
                setattr(self, f"btn_h{lvl}", btn)
                level_row.addWidget(btn)
            level_row.addStretch()
            vl.addLayout(level_row)

            self.header = QTextEdit()
            self.header.setPlaceholderText(f"Heading {level}")
            self.header.setMaximumHeight(50)
            self._apply_level_style()
            vl.addWidget(self.header)

        def _set_level(self, level):
            self.level = level
            for lvl in (1, 2, 3):
                getattr(self, f"btn_h{lvl}").setChecked(lvl == level)
            self._apply_level_style()

        def _apply_level_style(self):
            _, style, _ = self.LEVEL_STYLES[self.level]
            self.header.setStyleSheet(
                f"background-color: #0f172a; border: 1px solid #334155; border-radius: 4px; padding: 8px; {style}"
            )
            self.header.setPlaceholderText(f"Heading {self.level}")

    class SectionBodyWidget(QWidget):
        """A standalone section body/paragraph entry."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            vl = QVBoxLayout(self)
            vl.setContentsMargins(0, 0, 0, 0)
            vl.setSpacing(5)

            self.body = QTextEdit()
            self.body.setPlaceholderText("Section Content")
            self.body.setMinimumHeight(60)
            self.body.setMaximumHeight(500)
            self.body.setStyleSheet("background-color: #0f172a; color: white; border: 1px solid #334155; border-radius: 4px; padding: 10px; font-size: 14px;")
            self.body.textChanged.connect(self._auto_adjust_height)
            vl.addWidget(self.body)
        
        def _auto_adjust_height(self):
            """Auto-adjust height based on content"""
            doc_height = self.body.document().size().height()
            new_height = min(max(int(doc_height) + 20, 60), 500)
            self.body.setFixedHeight(new_height)

    class CodeBlockWidget(QWidget):
        """Monospaced code block with dark background."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            vl = QVBoxLayout(self)
            vl.setContentsMargins(0, 0, 0, 0)
            vl.setSpacing(2)
            label = QLabel("{ } Code Block")
            label.setStyleSheet("color: #22d3ee; font-size: 10px; font-weight: bold; padding: 2px 6px; background-color: #0e7490; border-radius: 3px;")
            label.setFixedHeight(18)
            vl.addWidget(label)
            self.code = QTextEdit()
            self.code.setPlaceholderText("Enter code here...")
            self.code.setMinimumHeight(80)
            self.code.setMaximumHeight(400)
            self.code.setStyleSheet("""
                QTextEdit {
                    background-color: #0d1117;
                    color: #22d3ee;
                    border: 1px solid #0e7490;
                    border-radius: 4px;
                    padding: 10px;
                    font-family: Consolas, 'Courier New', monospace;
                    font-size: 13px;
                }
            """)
            self.code.textChanged.connect(self._auto_adjust)
            vl.addWidget(self.code)

        def _auto_adjust(self):
            h = self.code.document().size().height()
            self.code.setFixedHeight(min(max(int(h) + 20, 80), 400))

    class QuoteBlockWidget(QWidget):
        """Blockquote widget with left border accent."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            hl = QHBoxLayout(self)
            hl.setContentsMargins(0, 0, 0, 0)
            hl.setSpacing(0)
            # Left accent bar
            bar = QFrame()
            bar.setFixedWidth(4)
            bar.setStyleSheet("background-color: #6366f1; border-radius: 2px;")
            hl.addWidget(bar)
            self.quote = QTextEdit()
            self.quote.setPlaceholderText("Enter quote or callout text...")
            self.quote.setMinimumHeight(60)
            self.quote.setMaximumHeight(300)
            self.quote.setStyleSheet("""
                QTextEdit {
                    background-color: #1e1b4b;
                    color: #a5b4fc;
                    border: none;
                    border-radius: 0 4px 4px 0;
                    padding: 10px 14px;
                    font-size: 14px;
                    font-style: italic;
                }
            """)
            self.quote.textChanged.connect(self._auto_adjust)
            hl.addWidget(self.quote)

        def _auto_adjust(self):
            h = self.quote.document().size().height()
            self.quote.setFixedHeight(min(max(int(h) + 20, 60), 300))

    class PageBreakWidget(QWidget):
        """Visual page break indicator."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            hl = QHBoxLayout(self)
            hl.setContentsMargins(0, 4, 0, 4)
            left = QFrame(); left.setFrameShape(QFrame.HLine)
            left.setStyleSheet("background-color: #475569;")
            lbl = QLabel("  ⊞ Page Break  ")
            lbl.setStyleSheet("color: #475569; font-size: 11px; font-weight: bold; background-color: black; padding: 0 6px;")
            lbl.setAlignment(Qt.AlignCenter)
            lbl.setFixedWidth(110)
            right = QFrame(); right.setFrameShape(QFrame.HLine)
            right.setStyleSheet("background-color: #475569;")
            hl.addWidget(left)
            hl.addWidget(lbl)
            hl.addWidget(right)

    class DividerWidget(QWidget):
        """A simple horizontal line divider widget."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            hl = QHBoxLayout(self)
            hl.setContentsMargins(0, 0, 0, 0)
            line = QFrame()
            line.setFrameShape(QFrame.HLine)
            line.setFrameShadow(QFrame.Sunken)
            line.setStyleSheet("background-color: #64748b; height: 2px;")
            hl.addWidget(line)

    class SpaceWidget(QWidget):
        """A simple vertical spacer widget."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            hl = QHBoxLayout(self)
            hl.setContentsMargins(0, 0, 0, 0)
            lbl = QLabel("[Empty Line]")
            lbl.setStyleSheet("color: #64748b; font-style: italic; font-size: 12px;")
            lbl.setAlignment(Qt.AlignCenter)
            hl.addWidget(lbl)

    class BulletWidget(QWidget):
        def __init__(self, number, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            layout = QHBoxLayout(self)
            layout.setContentsMargins(0, 0, 0, 0)
            self.label = QLabel(f"{number}.")
            self.label.setStyleSheet("color: white; font-weight: bold; font-size: 14px;")
            self.text = QLineEdit()
            self.text.setPlaceholderText("Bullet text")
            self.text.setStyleSheet("color: white; background-color: #0f172a; border: 1px solid #334155; border-radius: 4px; padding: 8px; font-size: 14px;")
            layout.addWidget(self.label)
            layout.addWidget(self.text)

    class CheckboxWidget(QWidget):
        """Bullet-like widget with a checkbox."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            layout = QHBoxLayout(self)
            layout.setContentsMargins(0, 0, 0, 0)
            self.checkbox = QCheckBox()
            self.text = QLineEdit()
            self.text.setPlaceholderText("Checkbox text")
            self.text.setStyleSheet("color: white; background-color: #0f172a; border: 1px solid #334155; border-radius: 4px; padding: 8px; font-size: 14px;")
            layout.addWidget(self.checkbox)
            layout.addWidget(self.text)

    class TableWidget(QWidget):
        """Simple editable table with move/remove controls and helpers to adjust size."""
        def __init__(self, remove_callback=None, move_up_callback=None, move_down_callback=None):
            super().__init__()
            vl = QVBoxLayout(self)
            vl.setContentsMargins(0, 0, 0, 0)
            vl.setSpacing(5)
            
            # helper toolbar for resizing
            helper_layout = QHBoxLayout()
            add_row_btn = QPushButton("+ Row")
            add_row_btn.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 4px 8px; border-radius: 4px;")
            add_col_btn = QPushButton("+ Col")
            add_col_btn.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 4px 8px; border-radius: 4px;")
            # header row toggle
            self.header_checkbox = QCheckBox("Header row")
            self.header_checkbox.setStyleSheet("color: white;")
            # border style selector
            self.border_combo = QComboBox()
            self.border_combo.addItems(["None", "Single", "Double"])
            self.border_combo.setStyleSheet("color: white; background-color: #0f172a; border: 1px solid #334155; padding: 4px;")
            helper_layout.addWidget(add_row_btn)
            helper_layout.addWidget(add_col_btn)
            helper_layout.addWidget(self.header_checkbox)
            helper_layout.addWidget(self.border_combo)
            helper_layout.addStretch()
            vl.addLayout(helper_layout)
            
            # start with 2x2 table for users to modify
            self.table = QTableWidget(2, 2)
            self.table.setStyleSheet("""
                QTableWidget { 
                    background-color: #0f172a; 
                    color: white; 
                    gridline-color: #334155; 
                    border: 1px solid #334155; 
                } 
                QHeaderView::section { 
                    background-color: #1e293b; 
                    color: white; 
                    border: 1px solid #334155; 
                    padding: 4px; 
                }
            """)
            # user typing layout configuration
            self.table.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            self.table.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
            self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
            self.table.verticalHeader().setSectionResizeMode(QHeaderView.ResizeToContents)
            # make the table take minimal vertical space out of the box but grow
            self.table.setMinimumHeight(100)
            
            # populate with empty QTableWidgetItems so user can type
            for r in range(2):
                for c in range(2):
                    self.table.setItem(r, c, QTableWidgetItem(""))
            
            vl.addWidget(self.table)
            
            def adapt_height():
                # calculate height required for all rows plus headers and margins
                h = self.table.horizontalHeader().height()
                for r in range(self.table.rowCount()):
                    h += self.table.rowHeight(r)
                self.table.setFixedHeight(h + 20)
                
            self.table.itemChanged.connect(adapt_height)
            QTimer.singleShot(0, adapt_height)
            # connect helpers
            def add_row():
                r = self.table.rowCount()
                self.table.insertRow(r)
                for c in range(self.table.columnCount()):
                    self.table.setItem(r, c, QTableWidgetItem(""))
                adapt_height()
            def add_col():
                c = self.table.columnCount()
                self.table.insertColumn(c)
                for r in range(self.table.rowCount()):
                    self.table.setItem(r, c, QTableWidgetItem(""))
                adapt_height()
            add_row_btn.clicked.connect(add_row)
            add_col_btn.clicked.connect(add_col)

    def __init__(self, stack, next_edit_screen):
        super().__init__()
        self.stack = stack
        self.next_edit_screen = next_edit_screen
        self.bullet_count = 0
        # store reusable titles
        self.titles = []
        self._load_titles()
        self.init_ui()

    def init_ui(self):
        # Root layout is vertical (Title -> Toolbar -> Main Area -> Nav Buttons)
        root_layout = QVBoxLayout()
        root_layout.setContentsMargins(20, 20, 20, 20)
        root_layout.setSpacing(10)
        self.setStyleSheet("background-color: #090d16;")

        title = QLabel("Create New Document")
        title.setObjectName("Title")
        title.setAlignment(Qt.AlignCenter)
        title.setStyleSheet("color: white; font-size: 24px; font-weight: bold;")
        root_layout.addWidget(title)

        # ========== UNIFIED TOOLBAR (2 rows) ==========
        toolbar_container = QFrame()
        toolbar_container.setStyleSheet("background-color: #0f172a; border: 1px solid #334155; border-radius: 8px; padding: 6px;")
        toolbar_container_layout = QVBoxLayout(toolbar_container)
        toolbar_container_layout.setSpacing(4)
        toolbar_container_layout.setContentsMargins(6, 6, 6, 6)

        # --- Row 1: Undo/Redo | Font Family | Font Size | Font Color | Strikethrough ---
        toolbar_row1 = QHBoxLayout()
        toolbar_row1.setSpacing(6)

        # Undo / Redo
        undo_redo_group = QFrame()
        undo_redo_group.setStyleSheet("background-color: #1e293b; border-radius: 4px; padding: 2px;")
        undo_redo_layout = QHBoxLayout(undo_redo_group)
        undo_redo_layout.setContentsMargins(4, 2, 4, 2)
        undo_redo_layout.setSpacing(4)

        self.toolbar_undo_btn = QPushButton("↩ Undo")
        self.toolbar_undo_btn.setFixedHeight(28)
        self.toolbar_undo_btn.setToolTip("Undo")
        self.toolbar_undo_btn.clicked.connect(self.toolbar_undo)
        undo_redo_layout.addWidget(self.toolbar_undo_btn)

        self.toolbar_redo_btn = QPushButton("↪ Redo")
        self.toolbar_redo_btn.setFixedHeight(28)
        self.toolbar_redo_btn.setToolTip("Redo")
        self.toolbar_redo_btn.clicked.connect(self.toolbar_redo)
        undo_redo_layout.addWidget(self.toolbar_redo_btn)
        toolbar_row1.addWidget(undo_redo_group)

        sep_r1a = QFrame(); sep_r1a.setFrameShape(QFrame.VLine)
        sep_r1a.setStyleSheet("background-color: #334155;")
        toolbar_row1.addWidget(sep_r1a)

        # Font Family
        font_family_label = QLabel("Font:")
        font_family_label.setStyleSheet("color: #94a3b8; font-size: 11px;")
        toolbar_row1.addWidget(font_family_label)
        self.toolbar_font_family = QFontComboBox()
        self.toolbar_font_family.setFixedWidth(150)
        self.toolbar_font_family.setFixedHeight(28)
        self.toolbar_font_family.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 2px;")
        self.toolbar_font_family.currentFontChanged.connect(self.toolbar_change_font_family)
        toolbar_row1.addWidget(self.toolbar_font_family)

        sep_r1b = QFrame(); sep_r1b.setFrameShape(QFrame.VLine)
        sep_r1b.setStyleSheet("background-color: #334155;")
        toolbar_row1.addWidget(sep_r1b)

        # Font Size
        font_size_label = QLabel("Size:")
        font_size_label.setStyleSheet("color: #94a3b8; font-size: 11px;")
        toolbar_row1.addWidget(font_size_label)
        self.toolbar_font_size = QSpinBox()
        self.toolbar_font_size.setRange(6, 96)
        self.toolbar_font_size.setValue(11)
        self.toolbar_font_size.setFixedWidth(60)
        self.toolbar_font_size.setFixedHeight(28)
        self.toolbar_font_size.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 2px;")
        self.toolbar_font_size.valueChanged.connect(self.toolbar_change_font_size)
        toolbar_row1.addWidget(self.toolbar_font_size)

        sep_r1c = QFrame(); sep_r1c.setFrameShape(QFrame.VLine)
        sep_r1c.setStyleSheet("background-color: #334155;")
        toolbar_row1.addWidget(sep_r1c)

        # Font Color
        self.toolbar_color_btn = QPushButton("🎨 Color")
        self.toolbar_color_btn.setFixedHeight(28)
        self.toolbar_color_btn.setToolTip("Font Color")
        self.toolbar_color_btn.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; border-radius: 4px; padding: 2px 8px;")
        self.toolbar_color_btn.clicked.connect(self.toolbar_pick_color)
        toolbar_row1.addWidget(self.toolbar_color_btn)
        self._current_font_color = QColor("white")

        sep_r1d = QFrame(); sep_r1d.setFrameShape(QFrame.VLine)
        sep_r1d.setStyleSheet("background-color: #334155;")
        toolbar_row1.addWidget(sep_r1d)

        # Strikethrough
        self.toolbar_strike_btn = QPushButton("S̶")
        self.toolbar_strike_btn.setFixedSize(36, 28)
        self.toolbar_strike_btn.setCheckable(True)
        self.toolbar_strike_btn.setToolTip("Strikethrough")
        self.toolbar_strike_btn.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; border-radius: 4px; text-decoration: line-through;")
        self.toolbar_strike_btn.clicked.connect(self.toolbar_toggle_strikethrough)
        toolbar_row1.addWidget(self.toolbar_strike_btn)

        sep_r1e = QFrame(); sep_r1e.setFrameShape(QFrame.VLine)
        sep_r1e.setStyleSheet("background-color: #334155;")
        toolbar_row1.addWidget(sep_r1e)

        # Line Spacing
        spacing_label = QLabel("Spacing:")
        spacing_label.setStyleSheet("color: #94a3b8; font-size: 11px;")
        toolbar_row1.addWidget(spacing_label)
        self.toolbar_spacing_combo = QComboBox()
        self.toolbar_spacing_combo.addItems(["1.0", "1.15", "1.5", "2.0", "2.5", "3.0"])
        self.toolbar_spacing_combo.setCurrentText("1.15")
        self.toolbar_spacing_combo.setFixedWidth(65)
        self.toolbar_spacing_combo.setFixedHeight(28)
        self.toolbar_spacing_combo.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 2px;")
        self.toolbar_spacing_combo.currentTextChanged.connect(self.toolbar_change_line_spacing)
        toolbar_row1.addWidget(self.toolbar_spacing_combo)

        sep_r1f = QFrame(); sep_r1f.setFrameShape(QFrame.VLine)
        sep_r1f.setStyleSheet("background-color: #334155;")
        toolbar_row1.addWidget(sep_r1f)

        # Indent
        indent_label = QLabel("Indent:")
        indent_label.setStyleSheet("color: #94a3b8; font-size: 11px;")
        toolbar_row1.addWidget(indent_label)
        self.toolbar_indent_btn = QPushButton("→")
        self.toolbar_indent_btn.setFixedSize(32, 28)
        self.toolbar_indent_btn.setToolTip("Increase Indent")
        self.toolbar_indent_btn.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; border-radius: 4px;")
        self.toolbar_indent_btn.clicked.connect(self.toolbar_increase_indent)
        toolbar_row1.addWidget(self.toolbar_indent_btn)
        self.toolbar_outdent_btn = QPushButton("←")
        self.toolbar_outdent_btn.setFixedSize(32, 28)
        self.toolbar_outdent_btn.setToolTip("Decrease Indent")
        self.toolbar_outdent_btn.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; border-radius: 4px;")
        self.toolbar_outdent_btn.clicked.connect(self.toolbar_decrease_indent)
        toolbar_row1.addWidget(self.toolbar_outdent_btn)

        toolbar_row1.addStretch()
        toolbar_container_layout.addLayout(toolbar_row1)

        # --- Row 2: existing controls ---
        toolbar = QFrame()
        toolbar.setStyleSheet("background-color: transparent;")
        toolbar_layout = QHBoxLayout(toolbar)
        toolbar_layout.setSpacing(8)
        toolbar_layout.setContentsMargins(0, 0, 0, 0)
        
        # Move/Delete Controls
        move_group = QFrame()
        move_group.setStyleSheet("background-color: #1e293b; border-radius: 4px; padding: 4px;")
        move_layout = QHBoxLayout(move_group)
        move_layout.setContentsMargins(4, 4, 4, 4)
        move_layout.setSpacing(4)
        
        self.toolbar_up_btn = QPushButton("🔼")
        self.toolbar_up_btn.setFixedSize(32, 32)
        self.toolbar_up_btn.setToolTip("Move Up")
        self.toolbar_up_btn.clicked.connect(self.toolbar_move_up)
        move_layout.addWidget(self.toolbar_up_btn)
        
        self.toolbar_down_btn = QPushButton("🔽")
        self.toolbar_down_btn.setFixedSize(32, 32)
        self.toolbar_down_btn.setToolTip("Move Down")
        self.toolbar_down_btn.clicked.connect(self.toolbar_move_down)
        move_layout.addWidget(self.toolbar_down_btn)
        
        self.toolbar_delete_btn = QPushButton("✖")
        self.toolbar_delete_btn.setFixedSize(32, 32)
        self.toolbar_delete_btn.setToolTip("Delete")
        self.toolbar_delete_btn.setStyleSheet("color: #ef4444; font-weight: bold;")
        self.toolbar_delete_btn.clicked.connect(self.toolbar_delete)
        move_layout.addWidget(self.toolbar_delete_btn)
        
        toolbar_layout.addWidget(move_group)
        
        # Separator
        sep1 = QFrame()
        sep1.setFrameShape(QFrame.VLine)
        sep1.setStyleSheet("background-color: #334155; width: 1px;")
        toolbar_layout.addWidget(sep1)
        
        # Text Formatting Controls
        format_group = QFrame()
        format_group.setStyleSheet("background-color: #1e293b; border-radius: 4px; padding: 4px;")
        format_layout = QHBoxLayout(format_group)
        format_layout.setContentsMargins(4, 4, 4, 4)
        format_layout.setSpacing(4)
        
        self.toolbar_bold_btn = QPushButton("B")
        self.toolbar_bold_btn.setFixedSize(32, 32)
        self.toolbar_bold_btn.setCheckable(True)
        self.toolbar_bold_btn.setToolTip("Bold")
        self.toolbar_bold_btn.setStyleSheet("font-weight: bold;")
        self.toolbar_bold_btn.clicked.connect(self.toolbar_toggle_bold)
        format_layout.addWidget(self.toolbar_bold_btn)
        
        self.toolbar_italic_btn = QPushButton("I")
        self.toolbar_italic_btn.setFixedSize(32, 32)
        self.toolbar_italic_btn.setCheckable(True)
        self.toolbar_italic_btn.setToolTip("Italic")
        self.toolbar_italic_btn.setStyleSheet("font-style: italic;")
        self.toolbar_italic_btn.clicked.connect(self.toolbar_toggle_italic)
        format_layout.addWidget(self.toolbar_italic_btn)
        
        self.toolbar_underline_btn = QPushButton("U")
        self.toolbar_underline_btn.setFixedSize(32, 32)
        self.toolbar_underline_btn.setCheckable(True)
        self.toolbar_underline_btn.setToolTip("Underline")
        self.toolbar_underline_btn.setStyleSheet("text-decoration: underline;")
        self.toolbar_underline_btn.clicked.connect(self.toolbar_toggle_underline)
        format_layout.addWidget(self.toolbar_underline_btn)
        
        toolbar_layout.addWidget(format_group)
        
        # Alignment Controls
        align_group = QFrame()
        align_group.setStyleSheet("background-color: #1e293b; border-radius: 4px; padding: 4px;")
        align_layout = QHBoxLayout(align_group)
        align_layout.setContentsMargins(4, 4, 4, 4)
        align_layout.setSpacing(4)
        
        self.toolbar_left_btn = QPushButton("Left")
        self.toolbar_left_btn.setFixedSize(50, 32)
        self.toolbar_left_btn.setToolTip("Align Left")
        self.toolbar_left_btn.clicked.connect(lambda: self.toolbar_set_alignment(Qt.AlignLeft))
        align_layout.addWidget(self.toolbar_left_btn)
        
        self.toolbar_center_btn = QPushButton("Center")
        self.toolbar_center_btn.setFixedSize(60, 32)
        self.toolbar_center_btn.setToolTip("Align Center")
        self.toolbar_center_btn.clicked.connect(lambda: self.toolbar_set_alignment(Qt.AlignCenter))
        align_layout.addWidget(self.toolbar_center_btn)
        
        self.toolbar_right_btn = QPushButton("Right")
        self.toolbar_right_btn.setFixedSize(50, 32)
        self.toolbar_right_btn.setToolTip("Align Right")
        self.toolbar_right_btn.clicked.connect(lambda: self.toolbar_set_alignment(Qt.AlignRight))
        align_layout.addWidget(self.toolbar_right_btn)
        
        toolbar_layout.addWidget(align_group)
        
        # Style Dropdown
        self.toolbar_style_combo = QComboBox()
        self.toolbar_style_combo.addItems(["Normal", "Heading 1", "Heading 2"])
        self.toolbar_style_combo.setFixedWidth(120)
        self.toolbar_style_combo.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 4px;")
        self.toolbar_style_combo.currentIndexChanged.connect(self.toolbar_change_style)
        toolbar_layout.addWidget(self.toolbar_style_combo)
        
        # Separator
        sep2 = QFrame()
        sep2.setFrameShape(QFrame.VLine)
        sep2.setStyleSheet("background-color: #334155; width: 1px;")
        toolbar_layout.addWidget(sep2)
        
        # Convert Element Type
        convert_label = QLabel("Convert to:")
        convert_label.setStyleSheet("color: #94a3b8; font-size: 12px;")
        toolbar_layout.addWidget(convert_label)
        
        self.toolbar_convert_combo = QComboBox()
        self.toolbar_convert_combo.addItems(["Header", "Paragraph", "Bullet", "Checkbox"])
        self.toolbar_convert_combo.setFixedWidth(100)
        self.toolbar_convert_combo.setStyleSheet("background-color: #1e293b; color: white; border: 1px solid #334155; padding: 4px;")
        toolbar_layout.addWidget(self.toolbar_convert_combo)
        
        self.toolbar_convert_btn = QPushButton("Convert")
        self.toolbar_convert_btn.setFixedSize(70, 32)
        self.toolbar_convert_btn.setStyleSheet("background-color: #f59e0b; color: white; border-radius: 4px; font-weight: bold;")
        self.toolbar_convert_btn.clicked.connect(self.toolbar_convert_element)
        toolbar_layout.addWidget(self.toolbar_convert_btn)
        
        toolbar_layout.addStretch()
        self.toolbar_status = QLabel("No selection")
        self.toolbar_status.setStyleSheet("color: #94a3b8; font-size: 12px;")
        toolbar_layout.addWidget(self.toolbar_status)

        toolbar_container_layout.addWidget(toolbar)
        root_layout.addWidget(toolbar_container)
        
        # Store current selection
        self.current_selected_widget = None
        self.update_toolbar_state()

        # Horizontal area for Sidebar + Main Editor
        self.main_h_layout = QHBoxLayout()
        self.main_h_layout.setSpacing(10)

        # --- Sidebar ---
        self.side_panel = QFrame()
        self.side_panel.setStyleSheet("background-color: #111827; border-right: 1px solid #374151; border-radius: 8px;")
        self.side_panel.setFixedWidth(200)
        self.side_layout = QVBoxLayout(self.side_panel)
        self.side_layout.setAlignment(Qt.AlignTop)
        self.side_layout.setSpacing(10)
        self.side_layout.setContentsMargins(15, 20, 15, 20)

        side_label = QLabel("TOOLBOX")
        side_label.setStyleSheet("font-weight: bold; color: #64748b; font-size: 11px; letter-spacing: 1px; margin-bottom: 5px;")
        self.side_layout.addWidget(side_label)

        def create_tool_btn(text, icon_char, callback):
            btn = QPushButton(f"{icon_char}  {text}")
            btn.setStyleSheet("""
                QPushButton {
                    background-color: transparent;
                    color: #cbd5e1;
                    border: none;
                    text-align: left;
                    padding: 8px 12px;
                    border-radius: 6px;
                    font-size: 13px;
                }
                QPushButton:hover {
                    background-color: #1e293b;
                    color: white;
                }
            """)
            btn.clicked.connect(callback)
            return btn

        self.side_layout.addWidget(create_tool_btn("Add H1 Header", "H1", lambda: self.insert_header(1)))
        self.side_layout.addWidget(create_tool_btn("Add H2 Header", "H2", lambda: self.insert_header(2)))
        self.side_layout.addWidget(create_tool_btn("Add H3 Header", "H3", lambda: self.insert_header(3)))
        self.side_layout.addWidget(create_tool_btn("Add Paragraph", "📝", self.insert_section_body))
        self.side_layout.addWidget(create_tool_btn("Add Bullets", "🔢", self.insert_bullets))
        self.side_layout.addWidget(create_tool_btn("Add Checkbox", "☑", self.insert_checkbox))
        self.side_layout.addWidget(create_tool_btn("Add Table", "📊", self.insert_table))
        self.side_layout.addWidget(create_tool_btn("Add Code Block", "{ }", self.insert_code_block))
        self.side_layout.addWidget(create_tool_btn("Add Quote", "❝", self.insert_quote_block))
        self.side_layout.addWidget(create_tool_btn("Add Spacing", "📏", self.insert_space))
        self.side_layout.addWidget(create_tool_btn("Add Divider", "➖", self.insert_divider))
        self.side_layout.addWidget(create_tool_btn("Add Page Break", "⊞", self.insert_page_break))

        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setStyleSheet("background-color: #334155; height: 1px; margin: 5px 0;")
        self.side_layout.addWidget(line)

        # Upload Document Button
        self.btn_upload = QPushButton("📤  Upload Document")
        self.btn_upload.clicked.connect(self.upload_document)
        self.btn_upload.setStyleSheet("""
            QPushButton {
                background-color: #10b981;
                color: white;
                border: none;
                padding: 10px;
                border-radius: 6px;
                font-weight: bold;
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        self.side_layout.addWidget(self.btn_upload)

        self.btn_ph = QPushButton("🔗  Insert Placeholder")
        self.btn_ph.clicked.connect(self.insert_placeholder)
        self.btn_ph.setStyleSheet("""
            QPushButton {
                background-color: #2563eb;
                color: white;
                border: none;
                padding: 10px;
                border-radius: 6px;
                font-weight: bold;
                margin-top: 5px;
            }
            QPushButton:hover {
                background-color: #3b82f6;
            }
        """)
        self.side_layout.addWidget(self.btn_ph)

        self.side_layout.addStretch()
        self.main_h_layout.addWidget(self.side_panel)

        # Toggle Button
        self.toggle_btn = QPushButton("◀")
        self.toggle_btn.setFixedWidth(20)
        self.toggle_btn.clicked.connect(self.toggle_sidebar)
        self.toggle_btn.setStyleSheet("background-color: #0f172a; border: 1px solid #334155; color: #94a3b8; border-radius: 4px;")
        self.main_h_layout.addWidget(self.toggle_btn)

        # --- Main Editor Area ---
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(True)
        self.scroll.setStyleSheet("background-color: #090d16; border: none;")
        self.content_container = QWidget()
        self.content_container.setStyleSheet("background-color: transparent;")
        self.content_layout = QVBoxLayout(self.content_container)
        self.content_layout.setAlignment(Qt.AlignTop)
        self.content_layout.setContentsMargins(40, 20, 40, 40)
        self.scroll.setWidget(self.content_container)
        self.main_h_layout.addWidget(self.scroll)

        root_layout.addLayout(self.main_h_layout)

        # Lower Nav Buttons
        btn_layout = QHBoxLayout()
        self.clear_btn = QPushButton("🗑 Clear All")
        self.clear_btn.clicked.connect(self.clear_all)
        self.clear_btn.setStyleSheet("color: #ef4444; background: transparent; border: 1px solid #7f1d1d; border-radius: 6px; padding: 6px 12px;")
        btn_layout.addWidget(self.clear_btn)
        
        btn_layout.addStretch()

        self.cancel_btn = QPushButton("Cancel")
        self.cancel_btn.clicked.connect(self.go_back)
        self.cancel_btn.setStyleSheet("padding: 6px 15px;")
        btn_layout.addWidget(self.cancel_btn)

        self.save_btn = QPushButton("Save…")
        self.save_btn.clicked.connect(self.save_document)
        self.save_btn.setStyleSheet("padding: 6px 15px;")
        btn_layout.addWidget(self.save_btn)

        self.next_btn = QPushButton("Next ➡")
        self.next_btn.clicked.connect(self.proceed)
        self.next_btn.setProperty("class", "primary") # Inherits premium blue style from styles.py
        self.next_btn.setStyleSheet("padding: 6px 20px;")
        btn_layout.addWidget(self.next_btn)

        root_layout.addLayout(btn_layout)
        self.setLayout(root_layout)

    def toggle_sidebar(self):
        if self.side_panel.isVisible():
            self.side_panel.hide()
            self.toggle_btn.setText("▶")
        else:
            self.side_panel.show()
            self.toggle_btn.setText("◀")

    # ---------- helper insert methods ----------
    def insert_header(self, level=1):
        widget = self.HeaderWidget(
            level=level,
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.header.setFocus()

    def insert_section_body(self):
        widget = self.SectionBodyWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.body.setFocus()

    def insert_bullets(self):
        # add a BulletWidget with next number
        self.bullet_count += 1
        widget = self.BulletWidget(
            self.bullet_count,
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.text.setFocus()

    def insert_table(self):
        widget = self.TableWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.table.setFocus()

    def insert_checkbox(self):
        widget = self.CheckboxWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.text.setFocus()

    def insert_code_block(self):
        widget = self.CodeBlockWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.code.setFocus()

    def insert_quote_block(self):
        widget = self.QuoteBlockWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self._connect_widget_signals(widget)
        self.scroll.show()
        widget.quote.setFocus()

    def insert_page_break(self):
        widget = self.PageBreakWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self.scroll.show()

    def insert_space(self):
        widget = self.SpaceWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self.scroll.show()

    def insert_divider(self):
        widget = self.DividerWidget(
            remove_callback=self._remove_widget,
            move_up_callback=self._move_widget_up,
            move_down_callback=self._move_widget_down
        )
        self.content_layout.addWidget(widget)
        self.scroll.show()

    def insert_placeholder(self):
        # Collect previously used placeholders
        placeholders = set()
        # Look in existing widgets
        for i in range(self.content_layout.count()):
            w = self.content_layout.itemAt(i).widget()
            txt = ""
            if hasattr(w, 'body'): txt = w.body.toPlainText()
            elif hasattr(w, 'text'): txt = w.text.text()
            elif hasattr(w, 'table'):
                for r in range(w.table.rowCount()):
                    for c in range(w.table.columnCount()):
                        it = w.table.item(r, c)
                        if it: txt += it.text() + " "
            
            # Simple regex search for {{...}}
            import re
            matches = re.findall(r'\{\{(.*?)\}\}', txt)
            placeholders.update(matches)

        if placeholders:
            ph_list = sorted(list(placeholders))
            # Offer selection or new
            items = ["-- New Placeholder --"] + ph_list
            item, ok = QInputDialog.getItem(self, "Insert Placeholder", "Select or type new:", items, 0, True)
            if not ok or not item: return
            text = ""
            if item == "-- New Placeholder --":
                text, ok2 = QInputDialog.getText(self, "New Placeholder", "Name:")
                if not ok2 or not text: return
            else:
                text = item
        else:
            text, ok = QInputDialog.getText(self, "Insert Placeholder", "Name:")
            if not ok or not text: return

        ph = f"{{{{{text}}}}}"
        
        # Smart cursor alignment: find where the user is actually typing
        focus_widget = QApplication.focusWidget()
        
        if isinstance(focus_widget, QTextEdit):
            # Works for SectionBodyWidget and other rich editors
            focus_widget.insertPlainText(ph)
            return
        elif isinstance(focus_widget, QLineEdit):
            # Works for Bullets, Checkboxes, or Section titles
            focus_widget.insert(ph)
            return
        elif isinstance(focus_widget, QTableWidget):
            # If a table cell is selected
            cur = focus_widget.currentItem()
            if cur:
                cur.setText(cur.text() + ph)
            return

        # Fallback: if no focus, try the last added widget as before
        if self.content_layout.count():
            last = self.content_layout.itemAt(self.content_layout.count() - 1).widget()
            if hasattr(last, 'body'):
                last.body.insertPlainText(ph)
                return
            if hasattr(last, 'text'):
                last.text.insert(ph)
                return
            if hasattr(last, 'table') and last.table.currentItem():
                cur = last.table.currentItem()
                cur.setText(cur.text() + ph)
                return
        # fallback: add a new section body containing placeholder
        self.insert_section_body()
        last = self.content_layout.itemAt(self.content_layout.count() - 1).widget()
        if hasattr(last, 'body'):
            last.body.setPlainText(ph)

    def _register_title_from_widget(self, widget):
        try:
            txt = widget.title.currentText().strip()
        except Exception:
            txt = ""
        if txt and txt not in self.titles:
            self.titles.append(txt)
            self._save_titles()
            # add to all existing section widgets' comboboxes
            for i in range(self.content_layout.count()):
                w = self.content_layout.itemAt(i).widget()
                if isinstance(w, (self.SectionWidget, self.SectionTitleWidget)) and w is not widget:
                    w.title.addItem(txt)

    # ---------- persistence helpers ----------
    def _titles_path(self):
        # store list in workspace root
        return os.path.join(os.getcwd(), "titles.json")

    def _load_titles(self):
        try:
            path = self._titles_path()
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    self.titles = json.load(f)
        except Exception:
            self.titles = []

    def _save_titles(self):
        try:
            path = self._titles_path()
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.titles, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    # ---------- widget management ----------
    def _remove_widget(self, widget):
        # remove from layout and delete
        self.content_layout.removeWidget(widget)
        widget.setParent(None)
        widget.deleteLater()
        self._renumber_bullets()

    def _move_widget_up(self, widget):
        # find current index and move up by one
        idx = -1
        for i in range(self.content_layout.count()):
            if self.content_layout.itemAt(i).widget() is widget:
                idx = i
                break
        if idx > 0:
            # remove and re-insert one position up
            self.content_layout.removeWidget(widget)
            self.content_layout.insertWidget(idx - 1, widget)
            self._renumber_bullets()

    def _move_widget_down(self, widget):
        # find current index and move down by one
        idx = -1
        for i in range(self.content_layout.count()):
            if self.content_layout.itemAt(i).widget() is widget:
                idx = i
                break
        if idx >= 0 and idx < self.content_layout.count() - 1:
            # remove and re-insert one position down
            self.content_layout.removeWidget(widget)
            self.content_layout.insertWidget(idx + 1, widget)
            self._renumber_bullets()

    def _renumber_bullets(self):
        count = 0
        for i in range(self.content_layout.count()):
            w = self.content_layout.itemAt(i).widget()
            if isinstance(w, self.BulletWidget):
                count += 1
                w.label.setText(f"{count}.")
        self.bullet_count = count

    def clear_all(self):
        reply = QMessageBox.question(self, "Clear All", "Are you sure you want to clear the entire document?", QMessageBox.Yes | QMessageBox.No)
        if reply == QMessageBox.Yes:
            while self.content_layout.count():
                item = self.content_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            self._renumber_bullets()
            self.scroll.hide()

    def go_back(self):
        # simply return to startup screen (index 0 assumed)
        self.stack.setCurrentIndex(0)

    def _export_textedit_to_para(self, te, para):
        """Export a QTextEdit's content with full formatting into a docx paragraph."""
        from docx.shared import Pt, RGBColor as RC, Cm
        from docx.enum.text import WD_LINE_SPACING

        # Apply paragraph-level formatting (indent + line spacing)
        cursor = te.textCursor()
        cursor.movePosition(QTextCursor.Start)
        block_fmt = cursor.blockFormat()
        left_margin = block_fmt.leftMargin()
        if left_margin > 0:
            para.paragraph_format.left_indent = Cm(left_margin / 37.8)  # px to cm approx
        line_height = block_fmt.lineHeight()
        if line_height > 0:
            factor = line_height / 100.0
            para.paragraph_format.line_spacing = factor
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        # Character-level formatting
        cursor.movePosition(QTextCursor.Start)
        current_text = ""
        current_fmt = None

        def fmt_key(f):
            fg = f.foreground().color()
            return (
                f.font().bold(), f.font().italic(), f.font().underline(),
                f.font().strikeOut(), f.font().family(), int(f.font().pointSize()),
                fg.red(), fg.green(), fg.blue()
            )

        def commit(text, f):
            if not text: return
            run = para.add_run(text)
            run.bold = f.font().bold()
            run.italic = f.font().italic()
            run.underline = f.font().underline()
            run.font.strike = f.font().strikeOut()
            if f.font().family():
                run.font.name = f.font().family()
            if f.font().pointSize() > 0:
                run.font.size = Pt(f.font().pointSize())
            fg = f.foreground().color()
            if fg.isValid() and fg != QColor("black") and fg.name() != "#000000":
                run.font.color.rgb = RC(fg.red(), fg.green(), fg.blue())

        while not cursor.atEnd():
            cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor)
            char_fmt = cursor.charFormat()
            text = cursor.selectedText()
            if current_fmt is None:
                current_fmt = char_fmt; current_text = text
            elif fmt_key(current_fmt) == fmt_key(char_fmt):
                current_text += text
            else:
                commit(current_text, current_fmt)
                current_fmt = char_fmt; current_text = text
            cursor.clearSelection()
        commit(current_text, current_fmt)

    def build_doc(self):
        """Construct a python-docx Document from the current widgets."""
        doc = Document()
        for i in range(self.content_layout.count()):
            item = self.content_layout.itemAt(i)
            widget = item.widget()
            if isinstance(widget, self.SectionWidget):
                title = widget.title.currentText().strip()
                body = widget.body.toPlainText().strip()
                if title:
                    doc.add_paragraph(title, style='Heading 1')
                if body:
                    doc.add_paragraph(body)
            elif isinstance(widget, self.HeaderWidget):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                _, _, docx_style = self.HeaderWidget.LEVEL_STYLES[widget.level]
                if widget.header.toPlainText().strip() == "": continue
                para = doc.add_paragraph(style=docx_style)
                qt_align = widget.header.alignment()
                if qt_align == Qt.AlignCenter: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif qt_align == Qt.AlignRight: para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._export_textedit_to_para(widget.header, para)
            elif isinstance(widget, self.PageBreakWidget):
                doc.add_page_break()
            elif isinstance(widget, self.DividerWidget):
                p = doc.add_paragraph()
                p.add_run("________________________________________________________________").bold = True

            elif isinstance(widget, self.SectionBodyWidget):
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if widget.body.toPlainText().strip() == "": continue
                para = doc.add_paragraph()
                qt_align = widget.body.alignment()
                if qt_align == Qt.AlignCenter: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif qt_align == Qt.AlignRight: para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._export_textedit_to_para(widget.body, para)

            elif isinstance(widget, self.CodeBlockWidget):
                if widget.code.toPlainText().strip() == "": continue
                para = doc.add_paragraph(style='No Spacing')
                run = para.add_run(widget.code.toPlainText())
                run.font.name = 'Courier New'
                from docx.shared import Pt, RGBColor as RC
                run.font.size = Pt(10)
                run.font.color.rgb = RC(0x22, 0xD3, 0xEE)
            elif isinstance(widget, self.QuoteBlockWidget):
                if widget.quote.toPlainText().strip() == "": continue
                para = doc.add_paragraph(style='Quote') if 'Quote' in [s.name for s in doc.styles] else doc.add_paragraph()
                self._export_textedit_to_para(widget.quote, para)
            elif isinstance(widget, self.SpaceWidget):
                doc.add_paragraph("")
            elif isinstance(widget, self.BulletWidget):
                text = widget.text.text().strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')
            elif isinstance(widget, self.CheckboxWidget):
                text = widget.text.text().strip()
                if text:
                    para = doc.add_paragraph()
                    sym = "☑" if widget.checkbox.isChecked() else "☐"
                    para.add_run(f"{sym} {text}")
            elif isinstance(widget, self.TableWidget):
                # translate table contents into docx table
                rows = widget.table.rowCount()
                cols = widget.table.columnCount()
                if rows > 0 and cols > 0:
                    style = 'Table Grid' if widget.border_combo.currentText() != 'None' else None
                    tbl = doc.add_table(rows=rows, cols=cols, style=style)
                    for r in range(rows):
                        for c in range(cols):
                            item = widget.table.item(r, c)
                            if item:
                                cell = tbl.cell(r, c)
                                cell.text = item.text()
                                if r == 0 and widget.header_checkbox.isChecked():
                                    # bold header row
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
        return doc

    def proceed(self):
        # gather and advance without debug dialogs
        try:
            doc = self.build_doc()
            set_template("", doc)
            # prepare form screen
            try:
                self.next_edit_screen.build_form_fields()
            except Exception as e:
                QMessageBox.warning(self, "Debug", f"build_form_fields error: {e}")
            QMessageBox.information(self, "Next", "Opening questionnaire screen...")
            self.stack.setCurrentWidget(self.next_edit_screen)
        except Exception as e:
            QMessageBox.critical(self, "Error in proceed", str(e))

    def save_document(self):
        """Prompt the user for a path and save the current document."""
        doc = self.build_doc()
        path, _ = QFileDialog.getSaveFileName(self, "Save document", "", "Word files (*.docx)")
        if path:
            try:
                # ensure extension
                if not path.lower().endswith('.docx'):
                    path += '.docx'
                doc.save(path)
                QMessageBox.information(self, "Saved", f"Document saved to {path}")
            except Exception as e:
                QMessageBox.critical(self, "Save error", str(e))

    def upload_document(self):
        """Upload existing Word document and populate the editor"""
        filepath, _ = QFileDialog.getOpenFileName(self, "Upload Document", "", "Word Documents (*.docx)")
        if not filepath:
            return
        
        try:
            doc = Document(filepath)
            
            # Clear existing content
            while self.content_layout.count():
                item = self.content_layout.takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            self.bullet_count = 0
            
            # Parse document structure properly
            for element in doc.element.body:
                # Check if it's a paragraph
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if para:
                        text = para.text.strip()
                        
                        # Skip empty paragraphs (add spacing)
                        if not text:
                            self.insert_space()
                            continue
                        
                        # Check for divider (line of underscores or dashes)
                        if all(c in '_-═' for c in text):
                            self.insert_divider()
                            continue
                        
                        # Check for checkbox items
                        if text.startswith('☑') or text.startswith('☐'):
                            widget = self.CheckboxWidget(
                                remove_callback=self._remove_widget,
                                move_up_callback=self._move_widget_up,
                                move_down_callback=self._move_widget_down
                            )
                            widget.checkbox.setChecked(text.startswith('☑'))
                            widget.text.setText(text[2:].strip())
                            self.content_layout.addWidget(widget)
                            self._connect_widget_signals(widget)
                            continue
                        
                        # Check for headings
                        if para.style.name.startswith('Heading'):
                            level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                            level = max(1, min(3, level))
                            widget = self.HeaderWidget(
                                level=level,
                                remove_callback=self._remove_widget,
                                move_up_callback=self._move_widget_up,
                                move_down_callback=self._move_widget_down
                            )
                            
                            # Preserve formatting
                            cursor = widget.header.textCursor()
                            cursor.movePosition(QTextCursor.Start)
                            
                            for run in para.runs:
                                cursor.insertText(run.text)
                                cursor.movePosition(QTextCursor.Left, QTextCursor.KeepAnchor, len(run.text))
                                fmt = cursor.charFormat()
                                font = fmt.font()
                                if run.bold:
                                    font.setBold(True)
                                if run.italic:
                                    font.setItalic(True)
                                if run.underline:
                                    font.setUnderline(True)
                                fmt.setFont(font)
                                cursor.setCharFormat(fmt)
                                cursor.movePosition(QTextCursor.End)
                            
                            self.content_layout.addWidget(widget)
                            self._connect_widget_signals(widget)
                            continue
                        
                        # Check for bullet/numbered lists
                        if para.style.name in ['List Bullet', 'List Number', 'List Paragraph']:
                            self.bullet_count += 1
                            widget = self.BulletWidget(
                                self.bullet_count,
                                remove_callback=self._remove_widget,
                                move_up_callback=self._move_widget_up,
                                move_down_callback=self._move_widget_down
                            )
                            widget.text.setText(text)
                            self.content_layout.addWidget(widget)
                            self._connect_widget_signals(widget)
                            continue
                        
                        # Regular paragraph with formatting
                        widget = self.SectionBodyWidget(
                            remove_callback=self._remove_widget,
                            move_up_callback=self._move_widget_up,
                            move_down_callback=self._move_widget_down
                        )
                        
                        # Preserve formatting
                        cursor = widget.body.textCursor()
                        cursor.movePosition(QTextCursor.Start)
                        
                        for run in para.runs:
                            cursor.insertText(run.text)
                            # Apply formatting
                            cursor.movePosition(QTextCursor.Left, QTextCursor.KeepAnchor, len(run.text))
                            fmt = cursor.charFormat()
                            font = fmt.font()
                            if run.bold:
                                font.setBold(True)
                            if run.italic:
                                font.setItalic(True)
                            if run.underline:
                                font.setUnderline(True)
                            fmt.setFont(font)
                            cursor.setCharFormat(fmt)
                            cursor.movePosition(QTextCursor.End)
                        
                        # Set alignment
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                            widget.body.setAlignment(Qt.AlignCenter)
                        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                            widget.body.setAlignment(Qt.AlignRight)
                        else:
                            widget.body.setAlignment(Qt.AlignLeft)
                        
                        self.content_layout.addWidget(widget)
                        self._connect_widget_signals(widget)
                
                # Check if it's a table
                elif element.tag.endswith('tbl'):
                    # Find the corresponding table object
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        widget = self.TableWidget(
                            remove_callback=self._remove_widget,
                            move_up_callback=self._move_widget_up,
                            move_down_callback=self._move_widget_down
                        )
                        
                        # Resize table to match
                        rows = len(table.rows)
                        cols = len(table.columns)
                        widget.table.setRowCount(rows)
                        widget.table.setColumnCount(cols)
                        
                        # Populate cells
                        for r, row in enumerate(table.rows):
                            for c, cell in enumerate(row.cells):
                                item = QTableWidgetItem(cell.text)
                                widget.table.setItem(r, c, item)
                        
                        # Check if first row looks like header (bold text)
                        if rows > 0:
                            first_row_bold = False
                            for cell in table.rows[0].cells:
                                for para in cell.paragraphs:
                                    for run in para.runs:
                                        if run.bold:
                                            first_row_bold = True
                                            break
                            widget.header_checkbox.setChecked(first_row_bold)
                        
                        self.content_layout.addWidget(widget)
                        self._connect_widget_signals(widget)
            
            self.scroll.show()
            QMessageBox.information(self, "Success", 
                f"Document uploaded successfully!\n\n"
                f"Detected elements:\n"
                f"- Titles/Headings\n"
                f"- Paragraphs with formatting\n"
                f"- Bullet lists\n"
                f"- Tables\n"
                f"- Checkboxes\n"
                f"- Dividers\n\n"
                f"You can now edit and save.")
            
        except Exception as e:
            import traceback
            QMessageBox.critical(self, "Error", f"Failed to upload document:\n{e}\n\n{traceback.format_exc()}")

    # ========== UNIFIED TOOLBAR METHODS ==========
    
    def _connect_widget_signals(self, widget):
        """Connect widget focus signals to update toolbar"""
        if hasattr(widget, 'body') and isinstance(widget.body, QTextEdit):
            widget.body.installEventFilter(self)
        elif hasattr(widget, 'header') and isinstance(widget.header, QTextEdit):
            widget.header.installEventFilter(self)
        elif hasattr(widget, 'code') and isinstance(widget.code, QTextEdit):
            widget.code.installEventFilter(self)
        elif hasattr(widget, 'quote') and isinstance(widget.quote, QTextEdit):
            widget.quote.installEventFilter(self)
        elif hasattr(widget, 'text') and isinstance(widget.text, QLineEdit):
            widget.text.installEventFilter(self)
    
    def eventFilter(self, obj, event):
        """Track focus changes to update toolbar"""
        if event.type() == QEvent.FocusIn:
            # Find the parent widget
            parent = obj.parent()
            while parent and not isinstance(parent, (self.SectionBodyWidget, self.HeaderWidget, self.BulletWidget, self.CheckboxWidget, self.CodeBlockWidget, self.QuoteBlockWidget)):
                parent = parent.parent()
            if parent:
                self.current_selected_widget = parent
                self.update_toolbar_state()
        return super().eventFilter(obj, event)
    
    def update_toolbar_state(self):
        """Update toolbar buttons based on current selection"""
        rich_controls = [
            self.toolbar_bold_btn, self.toolbar_italic_btn, self.toolbar_underline_btn,
            self.toolbar_left_btn, self.toolbar_center_btn, self.toolbar_right_btn,
            self.toolbar_font_family, self.toolbar_font_size, self.toolbar_color_btn,
            self.toolbar_strike_btn, self.toolbar_undo_btn, self.toolbar_redo_btn,
            self.toolbar_spacing_combo, self.toolbar_indent_btn, self.toolbar_outdent_btn,
        ]
        if not self.current_selected_widget:
            self.toolbar_status.setText("No selection")
            for c in rich_controls: c.setEnabled(False)
            self.toolbar_style_combo.setEnabled(False)
            self.toolbar_up_btn.setEnabled(False)
            self.toolbar_down_btn.setEnabled(False)
            self.toolbar_delete_btn.setEnabled(False)
            self.toolbar_convert_combo.setEnabled(False)
            self.toolbar_convert_btn.setEnabled(False)
            return

        self.toolbar_up_btn.setEnabled(True)
        self.toolbar_down_btn.setEnabled(True)
        self.toolbar_delete_btn.setEnabled(True)

        is_rich = isinstance(self.current_selected_widget, (self.SectionBodyWidget, self.HeaderWidget, self.CodeBlockWidget, self.QuoteBlockWidget))
        for c in rich_controls: c.setEnabled(is_rich)
        self.toolbar_convert_combo.setEnabled(True)
        self.toolbar_convert_btn.setEnabled(True)

        if isinstance(self.current_selected_widget, self.SectionBodyWidget):
            self.toolbar_status.setText("Paragraph selected")
            self.toolbar_style_combo.setEnabled(True)
            font = self.current_selected_widget.body.currentFont()
        elif isinstance(self.current_selected_widget, self.HeaderWidget):
            self.toolbar_status.setText("Header selected")
            self.toolbar_style_combo.setEnabled(False)
            font = self.current_selected_widget.header.currentFont()
        elif isinstance(self.current_selected_widget, self.CodeBlockWidget):
            self.toolbar_status.setText("Code Block selected")
            self.toolbar_style_combo.setEnabled(False)
            font = self.current_selected_widget.code.currentFont()
        elif isinstance(self.current_selected_widget, self.QuoteBlockWidget):
            self.toolbar_status.setText("Quote selected")
            self.toolbar_style_combo.setEnabled(False)
            font = self.current_selected_widget.quote.currentFont()
        elif isinstance(self.current_selected_widget, self.BulletWidget):
            self.toolbar_status.setText("Bullet selected")
            self.toolbar_style_combo.setEnabled(False)
            return
        elif isinstance(self.current_selected_widget, self.CheckboxWidget):
            self.toolbar_status.setText("Checkbox selected")
            self.toolbar_style_combo.setEnabled(False)
            return
        else:
            self.toolbar_status.setText(f"{type(self.current_selected_widget).__name__} selected")
            self.toolbar_style_combo.setEnabled(False)
            return

        # Sync rich text button states
        self.toolbar_bold_btn.setChecked(font.bold())
        self.toolbar_italic_btn.setChecked(font.italic())
        self.toolbar_underline_btn.setChecked(font.underline())
        self.toolbar_strike_btn.setChecked(font.strikeOut())
        self.toolbar_font_size.blockSignals(True)
        self.toolbar_font_size.setValue(int(font.pointSize()) if font.pointSize() > 0 else 11)
        self.toolbar_font_size.blockSignals(False)
    
    def toolbar_move_up(self):
        if self.current_selected_widget:
            self._move_widget_up(self.current_selected_widget)
    
    def toolbar_move_down(self):
        if self.current_selected_widget:
            self._move_widget_down(self.current_selected_widget)
    
    def toolbar_delete(self):
        if self.current_selected_widget:
            self._remove_widget(self.current_selected_widget)
            self.current_selected_widget = None
            self.update_toolbar_state()
    
    def toolbar_undo(self):
        w = self.current_selected_widget
        if isinstance(w, self.SectionBodyWidget): w.body.undo()
        elif isinstance(w, self.HeaderWidget): w.header.undo()

    def toolbar_redo(self):
        w = self.current_selected_widget
        if isinstance(w, self.SectionBodyWidget): w.body.redo()
        elif isinstance(w, self.HeaderWidget): w.header.redo()

    def _active_textedit(self):
        w = self.current_selected_widget
        if isinstance(w, self.SectionBodyWidget): return w.body
        if isinstance(w, self.HeaderWidget): return w.header
        if isinstance(w, self.CodeBlockWidget): return w.code
        if isinstance(w, self.QuoteBlockWidget): return w.quote
        return None

    def toolbar_change_font_family(self, font):
        te = self._active_textedit()
        if not te: return
        fmt = QTextCharFormat()
        fmt.setFontFamily(font.family())
        self._apply_format(te, fmt)

    def toolbar_change_font_size(self, size):
        te = self._active_textedit()
        if not te: return
        fmt = QTextCharFormat()
        fmt.setFontPointSize(size)
        self._apply_format(te, fmt)

    def toolbar_pick_color(self):
        te = self._active_textedit()
        if not te: return
        color = QColorDialog.getColor(self._current_font_color, self, "Pick Font Color")
        if color.isValid():
            self._current_font_color = color
            self.toolbar_color_btn.setStyleSheet(
                f"background-color: #1e293b; color: {color.name()}; border: 2px solid {color.name()}; border-radius: 4px; padding: 2px 8px; font-weight: bold;"
            )
            fmt = QTextCharFormat()
            fmt.setForeground(color)
            self._apply_format(te, fmt)

    def toolbar_toggle_strikethrough(self):
        te = self._active_textedit()
        if not te: return
        fmt = QTextCharFormat()
        fmt.setFontStrikeOut(self.toolbar_strike_btn.isChecked())
        self._apply_format(te, fmt)

    def _apply_format(self, te, fmt):
        cursor = te.textCursor()
        if cursor.hasSelection():
            cursor.mergeCharFormat(fmt)
        else:
            # Select all text to apply format to entire content
            cursor.select(QTextCursor.Document)
            cursor.mergeCharFormat(fmt)
            # Also set as current char format for future typing
            te.mergeCurrentCharFormat(fmt)
        te.setFocus()

    def toolbar_change_line_spacing(self, value):
        te = self._active_textedit()
        if not te: return
        try:
            factor = float(value)
        except ValueError:
            return
        cursor = te.textCursor()
        # Apply to whole document if no selection
        if not cursor.hasSelection():
            cursor.select(QTextCursor.Document)
        block_fmt = cursor.blockFormat()
        block_fmt.setLineHeight(factor * 100, 1)  # 1 = ProportionalHeight
        cursor.setBlockFormat(block_fmt)
        te.setFocus()

    def toolbar_increase_indent(self):
        te = self._active_textedit()
        if not te: return
        cursor = te.textCursor()
        block_fmt = cursor.blockFormat()
        block_fmt.setLeftMargin(block_fmt.leftMargin() + 20)
        cursor.setBlockFormat(block_fmt)
        te.setFocus()

    def toolbar_decrease_indent(self):
        te = self._active_textedit()
        if not te: return
        cursor = te.textCursor()
        block_fmt = cursor.blockFormat()
        new_margin = max(0, block_fmt.leftMargin() - 20)
        block_fmt.setLeftMargin(new_margin)
        cursor.setBlockFormat(block_fmt)
        te.setFocus()

    def toolbar_toggle_bold(self):
        if isinstance(self.current_selected_widget, self.SectionBodyWidget):
            font = self.current_selected_widget.body.currentFont()
            font.setBold(self.toolbar_bold_btn.isChecked())
            self.current_selected_widget.body.setCurrentFont(font)
            self.current_selected_widget.body.setFocus()
        elif isinstance(self.current_selected_widget, self.HeaderWidget):
            font = self.current_selected_widget.header.currentFont()
            font.setBold(self.toolbar_bold_btn.isChecked())
            self.current_selected_widget.header.setCurrentFont(font)
            self.current_selected_widget.header.setFocus()
    
    def toolbar_toggle_italic(self):
        if isinstance(self.current_selected_widget, self.SectionBodyWidget):
            font = self.current_selected_widget.body.currentFont()
            font.setItalic(self.toolbar_italic_btn.isChecked())
            self.current_selected_widget.body.setCurrentFont(font)
            self.current_selected_widget.body.setFocus()
        elif isinstance(self.current_selected_widget, self.HeaderWidget):
            font = self.current_selected_widget.header.currentFont()
            font.setItalic(self.toolbar_italic_btn.isChecked())
            self.current_selected_widget.header.setCurrentFont(font)
            self.current_selected_widget.header.setFocus()
    
    def toolbar_toggle_underline(self):
        if isinstance(self.current_selected_widget, self.SectionBodyWidget):
            font = self.current_selected_widget.body.currentFont()
            font.setUnderline(self.toolbar_underline_btn.isChecked())
            self.current_selected_widget.body.setCurrentFont(font)
            self.current_selected_widget.body.setFocus()
        elif isinstance(self.current_selected_widget, self.HeaderWidget):
            font = self.current_selected_widget.header.currentFont()
            font.setUnderline(self.toolbar_underline_btn.isChecked())
            self.current_selected_widget.header.setCurrentFont(font)
            self.current_selected_widget.header.setFocus()
    
    def toolbar_set_alignment(self, alignment):
        if isinstance(self.current_selected_widget, self.SectionBodyWidget):
            self.current_selected_widget.body.setAlignment(alignment)
            self.current_selected_widget.body.setFocus()
        elif isinstance(self.current_selected_widget, self.HeaderWidget):
            self.current_selected_widget.header.setAlignment(alignment)
            self.current_selected_widget.header.setFocus()
    
    def toolbar_change_style(self, index):
        if isinstance(self.current_selected_widget, self.SectionBodyWidget):
            cursor = self.current_selected_widget.body.textCursor()
            if not cursor.hasSelection():
                cursor.select(QTextCursor.BlockUnderCursor)
            
            fmt = cursor.blockCharFormat()
            if index == 0:  # Normal
                fmt.setFontPointSize(10.5)
                fmt.setFontWeight(400)
            elif index == 1:  # Heading 1
                fmt.setFontPointSize(18)
                fmt.setFontWeight(700)
            elif index == 2:  # Heading 2
                fmt.setFontPointSize(14)
                fmt.setFontWeight(700)
            
            cursor.setBlockCharFormat(fmt)
            self.current_selected_widget.body.setFocus()
    
    def toolbar_convert_element(self):
        """Convert current element to another type"""
        if not self.current_selected_widget:
            QMessageBox.warning(self, "No Selection", "Please select an element to convert.")
            return
        
        # Get current index
        current_idx = -1
        for i in range(self.content_layout.count()):
            if self.content_layout.itemAt(i).widget() is self.current_selected_widget:
                current_idx = i
                break
        
        if current_idx == -1:
            return
        
        # Get content from current widget
        content = ""
        if isinstance(self.current_selected_widget, self.HeaderWidget):
            content = self.current_selected_widget.header.toPlainText()
        elif isinstance(self.current_selected_widget, self.SectionBodyWidget):
            content = self.current_selected_widget.body.toPlainText()
        elif isinstance(self.current_selected_widget, self.BulletWidget):
            content = self.current_selected_widget.text.text()
        elif isinstance(self.current_selected_widget, self.CheckboxWidget):
            content = self.current_selected_widget.text.text()
        else:
            QMessageBox.warning(self, "Cannot Convert", "This element type cannot be converted.")
            return
        
        # Create new widget based on selection
        convert_type = self.toolbar_convert_combo.currentText()
        new_widget = None
        
        if convert_type == "Header":
            new_widget = self.HeaderWidget(
                level=1,
                remove_callback=self._remove_widget,
                move_up_callback=self._move_widget_up,
                move_down_callback=self._move_widget_down
            )
            new_widget.header.setPlainText(content)
        
        elif convert_type == "Paragraph":
            new_widget = self.SectionBodyWidget(
                remove_callback=self._remove_widget,
                move_up_callback=self._move_widget_up,
                move_down_callback=self._move_widget_down
            )
            new_widget.body.setPlainText(content)
        
        elif convert_type == "Bullet":
            self.bullet_count += 1
            new_widget = self.BulletWidget(
                self.bullet_count,
                remove_callback=self._remove_widget,
                move_up_callback=self._move_widget_up,
                move_down_callback=self._move_widget_down
            )
            new_widget.text.setText(content)
        
        elif convert_type == "Checkbox":
            new_widget = self.CheckboxWidget(
                remove_callback=self._remove_widget,
                move_up_callback=self._move_widget_up,
                move_down_callback=self._move_widget_down
            )
            new_widget.text.setText(content)
        
        if new_widget:
            # Remove old widget
            self.content_layout.removeWidget(self.current_selected_widget)
            self.current_selected_widget.deleteLater()
            
            # Insert new widget at same position
            self.content_layout.insertWidget(current_idx, new_widget)
            self._connect_widget_signals(new_widget)
            
            # Update selection
            self.current_selected_widget = new_widget
            self.update_toolbar_state()
            
            # Renumber bullets if needed
            self._renumber_bullets()
            
            QMessageBox.information(self, "Converted", f"Element converted to {convert_type}!")

