from PySide6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, 
    QLineEdit, QTextEdit, QFileDialog, QMessageBox, QScrollArea,
    QFormLayout, QSplitter
)
from PySide6.QtCore import Qt
from docx import Document
import re

class SimpleDocumentEditor(QWidget):
    """Simplified document editor - UI to Word and Word to UI"""
    
    def __init__(self):
        super().__init__()
        self.doc = None
        self.fields = {}
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Header
        header = QHBoxLayout()
        title = QLabel("Simple Document Editor")
        title.setStyleSheet("font-size: 20px; font-weight: bold; color: #3b82f6;")
        header.addWidget(title)
        header.addStretch()
        
        # Upload/Create buttons
        self.upload_btn = QPushButton("📤 Upload Word Doc")
        self.upload_btn.clicked.connect(self.upload_word)
        header.addWidget(self.upload_btn)
        
        self.new_btn = QPushButton("📄 New Template")
        self.new_btn.clicked.connect(self.load_template)
        header.addWidget(self.new_btn)
        
        layout.addLayout(header)
        
        # Splitter for form and preview
        splitter = QSplitter(Qt.Horizontal)
        
        # Left: Form fields
        form_widget = QWidget()
        form_layout = QVBoxLayout(form_widget)
        
        form_label = QLabel("Fields")
        form_label.setStyleSheet("font-weight: bold; font-size: 16px;")
        form_layout.addWidget(form_label)
        
        self.form_scroll = QScrollArea()
        self.form_scroll.setWidgetResizable(True)
        self.form_container = QWidget()
        self.form_fields = QFormLayout(self.form_container)
        self.form_scroll.setWidget(self.form_container)
        form_layout.addWidget(self.form_scroll)
        
        splitter.addWidget(form_widget)
        
        # Right: Preview
        preview_widget = QWidget()
        preview_layout = QVBoxLayout(preview_widget)
        
        preview_label = QLabel("Preview")
        preview_label.setStyleSheet("font-weight: bold; font-size: 16px;")
        preview_layout.addWidget(preview_label)
        
        self.preview = QTextEdit()
        self.preview.setReadOnly(True)
        preview_layout.addWidget(self.preview)
        
        splitter.addWidget(preview_widget)
        splitter.setSizes([400, 600])
        
        layout.addWidget(splitter)
        
        # Bottom buttons
        bottom = QHBoxLayout()
        bottom.addStretch()
        
        self.save_btn = QPushButton("💾 Save as Word")
        self.save_btn.clicked.connect(self.save_word)
        self.save_btn.setEnabled(False)
        bottom.addWidget(self.save_btn)
        
        layout.addLayout(bottom)
    
    def load_template(self):
        """Load a Word template with placeholders"""
        path, _ = QFileDialog.getOpenFileName(self, "Select Template", "", "Word (*.docx)")
        if not path:
            return
        
        try:
            self.doc = Document(path)
            placeholders = self.extract_placeholders()
            self.build_form(placeholders)
            self.update_preview()
            self.save_btn.setEnabled(True)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load: {e}")
    
    def upload_word(self):
        """Upload existing Word doc and extract data to UI"""
        path, _ = QFileDialog.getOpenFileName(self, "Upload Document", "", "Word (*.docx)")
        if not path:
            return
        
        try:
            self.doc = Document(path)
            placeholders = self.extract_placeholders()
            self.build_form(placeholders)
            self.update_preview()
            self.save_btn.setEnabled(True)
            
            QMessageBox.information(self, "Success", "Document loaded! Edit fields and save.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to upload: {e}")
    
    def extract_placeholders(self):
        """Extract {{placeholder}} or <<placeholder>> from document"""
        placeholders = set()
        for para in self.doc.paragraphs:
            placeholders.update(re.findall(r'\{\{(\w+)\}\}', para.text))
            placeholders.update(re.findall(r'<<(\w+)>>', para.text))
        
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        placeholders.update(re.findall(r'\{\{(\w+)\}\}', para.text))
                        placeholders.update(re.findall(r'<<(\w+)>>', para.text))
        
        return sorted(placeholders)
    
    def build_form(self, placeholders):
        """Build form fields from placeholders"""
        # Clear existing
        while self.form_fields.count():
            item = self.form_fields.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        self.fields.clear()
        
        for placeholder in placeholders:
            field = QLineEdit()
            field.setPlaceholderText(f"Enter {placeholder}")
            field.textChanged.connect(self.update_preview)
            
            label = QLabel(f"{placeholder}:")
            label.setStyleSheet("font-weight: bold;")
            
            self.form_fields.addRow(label, field)
            self.fields[placeholder] = field
    
    def update_preview(self):
        """Update preview with current field values"""
        if not self.doc:
            return
        
        preview_text = []
        for para in self.doc.paragraphs:
            text = para.text
            for key, field in self.fields.items():
                val = field.text() or f"{{{{{key}}}}}"
                text = text.replace(f"{{{{{key}}}}}", val)
                text = text.replace(f"<<{key}>>", val)
            preview_text.append(text)
        
        self.preview.setPlainText("\n".join(preview_text))
    
    def save_word(self):
        """Save filled document as Word - Simple text replacement"""
        if not self.doc:
            return
        
        path, _ = QFileDialog.getSaveFileName(self, "Save Document", "", "Word (*.docx)")
        if not path:
            return
        
        try:
            # Simple text replacement in paragraphs
            for para in self.doc.paragraphs:
                for key, field in self.fields.items():
                    val = field.text()
                    if val:
                        for run in para.runs:
                            run.text = run.text.replace(f"{{{{{key}}}}}", val)
                            run.text = run.text.replace(f"<<{key}>>", val)
            
            # Simple text replacement in tables
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, field in self.fields.items():
                                val = field.text()
                                if val:
                                    for run in para.runs:
                                        run.text = run.text.replace(f"{{{{{key}}}}}", val)
                                        run.text = run.text.replace(f"<<{key}>>", val)
            
            self.doc.save(path)
            QMessageBox.information(self, "Success", f"Saved to:\n{path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save: {e}")
